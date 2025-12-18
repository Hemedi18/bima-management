from django.shortcuts import render, redirect, get_object_or_404
from .models import Student
from django.db.models import Count, Q
from .forms import StudentUpdateForm, StudentCreateForm
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponse
import csv
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches


# Create your views here.
@login_required
@user_passes_test(lambda u: u.is_superuser)
def student_dashboard(request):
    # Chuja wanafunzi kulingana na hali ya usajili
    filter_by = request.GET.get('filter')
    if filter_by == 'registered':
        students = Student.objects.filter(is_registered=True)
    elif filter_by == 'unregistered':
        students = Student.objects.filter(is_registered=False)
    else:
        students = Student.objects.all()

    # Takwimu zilizoboreshwa
    stats = Student.objects.aggregate(
        total_students=Count('id'),
        registered_count=Count('id', filter=Q(is_registered=True))
    )
    total_students = stats['total_students']
    registered_count = stats['registered_count']
    unregistered_count = total_students - registered_count

    context = {
        'students': students,
        'total_students': total_students,
        'registered_count': registered_count,
        'unregistered_count': unregistered_count,
        'filter_by': filter_by,
    }
    return render(request, 'students/student_dashboard.html', context)

def add_student(request):
    if request.method == 'POST':
        form = StudentCreateForm(request.POST)
        if form.is_valid():
            student = form.save(commit=False)
            student.is_registered = bool(student.phone_number)
            student.save()
            messages.success(request, f"Umefanikiwa kujisajili, {student.full_name}! Sasa unaweza kuingia ili kusasisha taarifa zako.")
            return redirect('student_login')
    else:
        form = StudentCreateForm()
    
    context = {
        'form': form
    }
    return render(request, 'students/add_student.html', context)

def student_login(request):
    if request.method == 'POST':
        reg_number = request.POST.get('q', '').strip()
        if reg_number:
            try:
                student = Student.objects.get(registration_number__iexact=reg_number)
                # Mwanafunzi amepatikana, mpeleke kwenye ukurasa wake
                messages.info(request, f"Karibu {student.full_name}. Hapa unaweza kuona na kubadilisha taarifa zako.")
                return redirect('student_update', pk=student.pk)
            except Student.DoesNotExist:
                messages.error(request, f"Mwanafunzi mwenye namba ya usajili '{reg_number}' hajapatikana. Tafadhali hakikisha umeandika kwa usahihi.")
        else:
            messages.error(request, "Tafadhali ingiza namba ya usajili.")
    
    return render(request, 'students/student_login.html')

def student_update(request, pk):
    student = get_object_or_404(Student, pk=pk)
    
    if request.method == 'POST':
        form = StudentUpdateForm(request.POST, instance=student)
        if form.is_valid():
            updated_student = form.save(commit=False)
            updated_student.is_registered = bool(updated_student.phone_number)
            if form.has_changed():
                updated_student.save()
                messages.success(request, 'Taarifa zako zimesasishwa kikamilifu!')
            else:
                messages.info(request, 'Hakuna mabadiliko yaliyofanywa.')
            return redirect('student_update', pk=student.pk)
    else:
        form = StudentUpdateForm(instance=student)

    context = {
        'form': form,
        'student': student
    }
    return render(request, 'students/student_update.html', context)

@login_required
@user_passes_test(lambda u: u.is_superuser)
def export_students(request):
    # The export is triggered by a GET request from the fetch API
    data_scope = request.GET.get('data_scope', 'all')
    file_format = request.GET.get('format', 'csv')

    # Base queryset
    students_qs = Student.objects.all().order_by('registration_number')

    # Filter based on data scope
    if data_scope == 'complete':
        # 'complete' means registered
        students_qs = students_qs.filter(is_registered=True)
    elif data_scope == 'incomplete':
        # 'incomplete' means not registered
        students_qs = students_qs.filter(is_registered=False)

    # Prepare data
    headers = ['S/N', 'Registration Number', 'Full Name', 'Phone Number']
    data = []
    for i, student in enumerate(students_qs, 1):
        data.append([
            i,
            student.registration_number or '',
            student.full_name or '',
            student.phone_number or ''
        ])

    # Generate file based on format
    if file_format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="students.csv"'
        writer = csv.writer(response)
        writer.writerow(headers)
        writer.writerows(data)
        return response

    elif file_format == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="students.pdf"'
        doc = SimpleDocTemplate(response, pagesize=letter)
        table_data = [headers] + data
        table = Table(table_data)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        table.setStyle(style)
        doc.build([table])
        return response

    elif file_format == 'docx':
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="students.docx"'
        document = Document()
        document.add_heading('Student Data', 0)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        document.save(response)
        return response

    return HttpResponse("Invalid format", status=400)
